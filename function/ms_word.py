import re
from datetime import datetime

from module.common import Common

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def make_replacer(data_dict):
    def replacer(match):
        key = match.group(1).lower()
        return str(data_dict.get(key, f"{{{key}}}"))
    return replacer


class MSword(Common):
    def __init__(self):
        self.logger = self.get_logger()

        self.sample_report_path = './sample/sample_healthcheck_report.docx'
        self.report_path = './result/[RockPLACE] 정기점검보고서_'

        self.graph_folder = './result/graphs/'

    def make_report(self, data):
        report = Document(self.sample_report_path)

        service_name = data["service_name"]
        # data["execute_time"] = data["execute_time"].strftime('%Y-%m-%d %H:%M:%S')

        # 왼쪽 정렬 파라미터 리스트 (여기 없으면 모두 가운데 정렬)
        left_aligned_param_list = [
            "engine_home", "data_home", "config_file", "error_log", "db_port",
            "replication_status", "innodb_engine_status", "error_log_content"
        ]

        ################## 데이터 삽입
        self.logger.info("input DB information & Graph on healthcheck report")

        for table in report.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()

                    if text == "{graph_datasize}":
                        self.insert_graph(cell.paragraphs[0], service_name)

                    elif "{" in text and "}" in text:
                        # 자리표시자 모두 치환
                        replaced_text = re.sub(r'{(.*?)}', make_replacer(data), text)

                        # 기존 텍스트 삭제 및 새 텍스트 삽입
                        cell.text = ""
                        para = cell.paragraphs[0]
                        run = para.add_run(replaced_text)
                        run.font.name = "Malgun Gothic"
                        run.font.size = Pt(7)

                        # 정렬 판단 (첫 키만 기준으로 정렬 적용)
                        first_key_match = re.search(r'{(.*?)}', text)
                        if first_key_match:
                            first_key = first_key_match.group(1).lower()
                            if first_key in left_aligned_param_list:
                                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            else:
                                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # report_name = self.report_path + service_name + "_" + datetime.strptime(data["execute_time"], '%Y-%m-%d %H:%M:%S').strftime('%y%m%d') + ".docx"
        report_name = self.report_path + service_name + "_" + data["execute_time"].strftime('%y%m%d') + ".docx"
        report.save(report_name)

    def insert_graph(self, paragraph, filename):
        """
        문단 내 플레이스홀더({graph_파일명})를 찾아 이미지를 삽입하는 함수
        """
        image_path = None
        image_extensions = [".png"]

        # 해당 파일명이 존재하는 이미지 찾기
        for ext in image_extensions:
            potential_path = f"./result/{filename}{ext}"
            try:
                with open(potential_path, "rb"):  # 파일 존재 확인
                    image_path = potential_path
                    break
            except FileNotFoundError:
                continue

        if image_path:
            paragraph.text = ""  # 기존 텍스트 제거
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(4))  # 이미지 크기 조정
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 가운데 정렬 적용
        else:
            self.logger.warning(f"이미지 파일을 찾을 수 없음: {filename}")



